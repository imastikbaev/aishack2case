from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Технический стек проекта AI-Завуч.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_heading(doc, text, size=14):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(7)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(21, 55, 78)
    return paragraph


def add_body(doc, text, bold_start=None):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.03
    if bold_start and text.startswith(bold_start):
        run = paragraph.add_run(bold_start)
        run.bold = True
        rest = text[len(bold_start):]
        paragraph.add_run(rest)
    else:
        paragraph.add_run(text)
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(9.5)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(1)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(9.2)
    return paragraph


def add_table(doc, rows):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["Слой", "Технологии", "Назначение в проекте"]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_text(cell, header, bold=True, color="FFFFFF")
        set_cell_shading(cell, "15374E")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    for row in table.rows:
        row.cells[0].width = Cm(3.1)
        row.cells[1].width = Cm(5.4)
        row.cells[2].width = Cm(8.5)
    return table


doc = Document()
section = doc.sections[0]
section.top_margin = Cm(1.25)
section.bottom_margin = Cm(1.2)
section.left_margin = Cm(1.35)
section.right_margin = Cm(1.35)

styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"].font.size = Pt(9.5)
styles["List Bullet"].font.name = "Calibri"
styles["List Bullet"].font.size = Pt(9.2)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(2)
run = title.add_run("AI-Завуч | Aqbobek Lyceum")
run.bold = True
run.font.name = "Calibri"
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(21, 55, 78)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(8)
run = subtitle.add_run("Технический стек и архитектура проекта")
run.font.name = "Calibri"
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(80, 80, 80)

add_heading(doc, "Кратко о проекте", 13)
add_body(
    doc,
    "AI-Завуч — интеллектуальная система управления школой: веб-дашборд для администрации, "
    "FastAPI backend, Telegram-бот для учителей и AI-модуль для разбора сообщений, RAG по нормативным "
    "документам, поиска замен, прогнозов и управленческих инсайтов."
)

add_heading(doc, "Архитектура", 13)
add_body(
    doc,
    "Frontend (React/Vite) отправляет запросы в REST API FastAPI. Backend хранит данные в SQLite через "
    "SQLAlchemy, выполняет бизнес-логику расписания и задач, обращается к Groq LLM при наличии ключа и "
    "использует smart mock-ответы без API-ключей. Telegram-бот принимает сообщения учителей, передает их "
    "в backend и показывает подтверждения/сводки."
)

stack_rows = [
    (
        "Frontend",
        "React 18, Vite 5, React Router, Tailwind CSS, Lucide React",
        "Одностраничное приложение с разделами: дашборд, расписание, задачи, посещаемость, инциденты, AI-ассистент, сотрудники.",
    ),
    (
        "Визуализация",
        "Recharts, date-fns, clsx",
        "Графики, карточки показателей, тепловые карты нагрузки, форматирование дат и условные UI-состояния.",
    ),
    (
        "Backend API",
        "Python, FastAPI 0.115, Uvicorn, Pydantic 2",
        "REST API для сотрудников, расписания, задач, посещаемости, инцидентов, уведомлений, AI-функций и dashboard-сводки.",
    ),
    (
        "Данные",
        "SQLAlchemy 2, SQLite",
        "Локальная база school.db с моделями Staff, ClassGroup, Room, Schedule, Task, Attendance, Incident, ChatMessage, Notification.",
    ),
    (
        "AI",
        "Groq SDK, llama-3.3-70b-versatile, RAG по приказам №76/110/130",
        "Разбор сообщений, Voice-to-Task, WhatsApp-парсинг поручений, рекомендации замен, прогноз посещаемости, анализ рисков и ответы по нормативным документам.",
    ),
    (
        "Мессенджеры",
        "python-telegram-bot 21.5, httpx, WhatsApp group workflow",
        "Telegram принимает отчеты по посещаемости и инцидентам. WhatsApp-сценарий отправляет структурированные задачи из голосовой команды директора в рабочую группу.",
    ),
    (
        "Интеграции",
        "Axios, CORS, python-dotenv, python-multipart, aiohttp/httpx",
        "Связь frontend-backend, настройка окружения, загрузка переменных, сетевые вызовы и подготовка к внешним сервисам.",
    ),
    (
        "Запуск",
        "start.sh, npm scripts, Python venv",
        "Одна команда поднимает backend на :8000, frontend на :5173, сидирует базу и запускает бота при наличии токена.",
    ),
]

add_heading(doc, "Технический стек", 13)
add_table(doc, stack_rows)

add_heading(doc, "Ключевые реализованные возможности", 13)
for item in [
    "Ежедневный dashboard: посещаемость, открытые инциденты, задачи, риски и AI-инсайты.",
    "Расписание: недельная сетка, расписание на сегодня, конфликты, тепловая карта нагрузки, умные замены.",
    "Задачи: канбан-доска, приоритеты, исполнители, сроки, создание задач из голосовой команды.",
    "WhatsApp-поручения: директор записывает голосовое сообщение на сайте, AI превращает его в структурированную задачу и готовит отправку в WhatsApp-группу.",
    "Посещаемость: отчеты по классам, история, прогноз на завтра, расчет порций в столовую.",
    "Инциденты: регистрация, статусы, категории, назначение ответственных и источник обращения.",
    "AI-ассистент: RAG по локальным нормативным файлам, симуляция Telegram-сообщений, анализ рисков персонала.",
]:
    add_bullet(doc, item)

add_heading(doc, "API и модули", 13)
add_body(
    doc,
    "Основные endpoint-группы: /api/staff, /api/schedule, /api/tasks, /api/attendance, /api/incidents, "
    "/api/messages, /api/ai, /api/notifications, /api/dashboard, /api/telegram. Для WhatsApp-сценария "
    "используется тот же AI-конвейер Voice-to-Task: аудио/текст поручения преобразуется в JSON-задачу "
    "с исполнителем, сроком, приоритетом и текстом для отправки в группу. Swagger UI доступен "
    "локально по адресу http://localhost:8000/docs."
)

add_heading(doc, "Что важно показать жюри", 13)
for item in [
    "Проект работает локально даже без Groq API key: AI-сценарии имеют fallback на smart mock-логику.",
    "Система закрывает полный школьный цикл: данные учителей, расписание, задачи, посещаемость, инциденты и уведомления.",
    "AI встроен в реальные процессы, а не вынесен отдельной игрушкой: он парсит сообщения, превращает голос директора в задачи, помогает с заменами и дает управленческие рекомендации.",
    "WhatsApp-группа используется как привычный канал доставки: сотрудники получают не сырой голос, а понятное поручение с ответственным, сроком и приоритетом.",
    "Архитектура модульная: frontend, backend, база данных, AI-сервис и Telegram-бот разделены и могут развиваться независимо.",
]:
    add_bullet(doc, item)

footer_section = doc.sections[0]
footer = footer_section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer.add_run("AI-Завуч | Технический стек проекта | для печати и защиты")
footer_run.font.name = "Calibri"
footer_run.font.size = Pt(8)
footer_run.font.color.rgb = RGBColor(120, 120, 120)

doc.save(OUT)
print(OUT)
